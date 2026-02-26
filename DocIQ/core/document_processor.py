"""
DocIQ — Document Processor
Handles PDF, DOCX, TXT with structure extraction and section detection.
"""

import io
import re
import logging
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Dict, Tuple

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────
# DATA MODELS
# ──────────────────────────────────────────────

@dataclass
class DocumentSection:
    """Represents a logical section within a document."""
    title: str
    level: int                  # heading depth (1=top-level, 2=sub, etc.)
    start_char: int
    end_char: int
    content: str
    section_index: int
    word_count: int = 0

    def __post_init__(self):
        self.word_count = len(self.content.split())


@dataclass
class ParsedDocument:
    """Complete parsed representation of a document."""
    filename: str
    format: str                             # pdf | docx | txt
    raw_text: str
    sections: List[DocumentSection]
    metadata: Dict = field(default_factory=dict)
    parse_time_ms: float = 0.0
    char_count: int = 0
    word_count: int = 0
    page_count: int = 0

    def __post_init__(self):
        self.char_count = len(self.raw_text)
        self.word_count = len(self.raw_text.split())


# ──────────────────────────────────────────────
# BASE PARSER
# ──────────────────────────────────────────────

class BaseParser:
    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        raise NotImplementedError

    def _detect_sections(self, text: str) -> List[DocumentSection]:
        """
        Regex-based section detection that works across document types.
        Identifies markdown headers, numbered sections, and ALL-CAPS titles.
        """
        patterns = [
            (r"^(#{1,6})\s+(.+)$", "markdown"),
            (r"^(\d+(?:\.\d+)*)\s+([A-Z].{2,80})$", "numbered"),
            (r"^([A-Z][A-Z\s\-]{4,60})$", "caps"),
            (r"^((?:Abstract|Introduction|Background|Methodology|Results|"
             r"Discussion|Conclusion|References|Appendix|Overview|Summary|"
             r"Analysis|Findings|Recommendations?|Executive Summary)\b.{0,60})$",
             "academic"),
        ]

        sections = []
        lines = text.split("\n")
        current_pos = 0
        section_starts = []   # (char_pos, title, level)

        for line in lines:
            stripped = line.strip()
            matched = False
            for pattern, ptype in patterns:
                m = re.match(pattern, stripped, re.IGNORECASE)
                if m and not matched:
                    if ptype == "markdown":
                        level = len(m.group(1))
                        title = m.group(2).strip()
                    elif ptype == "numbered":
                        level = m.group(1).count(".") + 1
                        title = m.group(2).strip()
                    else:
                        level = 1
                        title = stripped
                    section_starts.append((current_pos, title, level))
                    matched = True
                    break
            current_pos += len(line) + 1  # +1 for \n

        # Build section objects from detected boundaries
        for i, (start, title, level) in enumerate(section_starts):
            end = section_starts[i + 1][0] if i + 1 < len(section_starts) else len(text)
            content = text[start:end].strip()
            if len(content) > 30:  # skip trivially short sections
                sections.append(DocumentSection(
                    title=title,
                    level=level,
                    start_char=start,
                    end_char=end,
                    content=content,
                    section_index=i,
                ))

        # If no sections detected, treat entire document as one section
        if not sections:
            sections = [DocumentSection(
                title="Document Content",
                level=1,
                start_char=0,
                end_char=len(text),
                content=text.strip(),
                section_index=0,
            )]

        return sections

    def _clean_text(self, text: str) -> str:
        """Normalize whitespace, remove control characters, fix encoding artifacts."""
        # Remove null bytes and control chars (except newlines and tabs)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        # Normalize unicode dashes and quotes
        text = text.replace("\u2019", "'").replace("\u2018", "'")
        text = text.replace("\u201c", '"').replace("\u201d", '"')
        text = text.replace("\u2013", "-").replace("\u2014", "--")
        # Collapse excessive blank lines (max 2)
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Collapse excessive spaces (preserve indentation structure)
        text = re.sub(r"[ \t]{3,}", "  ", text)
        return text.strip()


# ──────────────────────────────────────────────
# PDF PARSER
# ──────────────────────────────────────────────

class PDFParser(BaseParser):
    """
    Two-layer PDF extraction:
    1. pdfplumber  — preferred (layout-aware, handles tables)
    2. PyPDF2      — fallback
    """

    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        t0 = time.time()
        text = ""
        page_count = 0
        metadata = {}

        # Layer 1: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                page_count = len(pdf.pages)
                metadata = pdf.metadata or {}
                pages_text = []
                for page in pdf.pages:
                    pt = page.extract_text(x_tolerance=3, y_tolerance=3)
                    if pt:
                        pages_text.append(pt)
                text = "\n\n".join(pages_text)
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}. Falling back to PyPDF2.")

        # Layer 2: PyPDF2 fallback
        if not text.strip():
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                page_count = len(reader.pages)
                pages_text = []
                for page in reader.pages:
                    pt = page.extract_text()
                    if pt:
                        pages_text.append(pt)
                text = "\n\n".join(pages_text)
                metadata = dict(reader.metadata) if reader.metadata else {}
            except Exception as e:
                logger.error(f"PyPDF2 also failed: {e}")
                text = "[PDF extraction failed. The document may be scanned or encrypted.]"

        text = self._clean_text(text)
        sections = self._detect_sections(text)

        return ParsedDocument(
            filename=filename,
            format="pdf",
            raw_text=text,
            sections=sections,
            metadata={k: str(v) for k, v in metadata.items() if v},
            parse_time_ms=(time.time() - t0) * 1000,
            page_count=page_count,
        )


# ──────────────────────────────────────────────
# DOCX PARSER
# ──────────────────────────────────────────────

class DOCXParser(BaseParser):
    """
    python-docx parser that preserves heading hierarchy and styles.
    """

    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        t0 = time.time()

        try:
            from docx import Document as DocxDocument
            from docx.oxml.ns import qn

            doc = DocxDocument(io.BytesIO(file_bytes))
            sections = []
            full_text_parts = []
            current_section_title = "Document Content"
            current_section_level = 1
            current_section_start = 0
            current_section_parts = []
            section_index = 0
            char_cursor = 0

            def flush_section():
                nonlocal section_index, current_section_start
                content = "\n".join(current_section_parts).strip()
                if len(content) > 30:
                    sections.append(DocumentSection(
                        title=current_section_title,
                        level=current_section_level,
                        start_char=current_section_start,
                        end_char=char_cursor,
                        content=content,
                        section_index=section_index,
                    ))
                    section_index += 1
                current_section_start = char_cursor

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else ""
                is_heading = "Heading" in style_name

                if is_heading:
                    # Extract heading level from style name
                    try:
                        level = int(re.search(r"\d+", style_name).group())
                    except Exception:
                        level = 1

                    flush_section()
                    current_section_title = text
                    current_section_level = level
                    current_section_parts = [text]
                else:
                    current_section_parts.append(text)
                    full_text_parts.append(text)
                    char_cursor += len(text) + 1

            flush_section()

            full_text = self._clean_text("\n".join(full_text_parts))

            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "author": core_props.author or "",
                "title": core_props.title or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
            }

            if not sections:
                sections = self._detect_sections(full_text)

            return ParsedDocument(
                filename=filename,
                format="docx",
                raw_text=full_text,
                sections=sections,
                metadata={k: v for k, v in metadata.items() if v},
                parse_time_ms=(time.time() - t0) * 1000,
            )

        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            return ParsedDocument(
                filename=filename,
                format="docx",
                raw_text=f"[DOCX extraction failed: {e}]",
                sections=[],
                parse_time_ms=(time.time() - t0) * 1000,
            )


# ──────────────────────────────────────────────
# TXT PARSER
# ──────────────────────────────────────────────

class TXTParser(BaseParser):
    """Plain text parser with encoding detection."""

    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        t0 = time.time()

        # Try UTF-8 first, then latin-1 fallback
        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                text = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            text = file_bytes.decode("utf-8", errors="replace")

        text = self._clean_text(text)
        sections = self._detect_sections(text)

        return ParsedDocument(
            filename=filename,
            format="txt",
            raw_text=text,
            sections=sections,
            parse_time_ms=(time.time() - t0) * 1000,
        )


# ──────────────────────────────────────────────
# DOCUMENT PROCESSOR (facade)
# ──────────────────────────────────────────────

class DocumentProcessor:
    """
    Unified document processing facade.
    Routes to the appropriate parser by file extension.
    """

    PARSERS = {
        ".pdf":  PDFParser,
        ".docx": DOCXParser,
        ".txt":  TXTParser,
        ".md":   TXTParser,
    }

    def process(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        ext = Path(filename).suffix.lower()
        parser_cls = self.PARSERS.get(ext)

        if not parser_cls:
            raise ValueError(f"Unsupported file format: {ext}. "
                             f"Supported: {list(self.PARSERS.keys())}")

        logger.info(f"Processing {filename} ({len(file_bytes)/1024:.1f} KB) with {parser_cls.__name__}")
        doc = parser_cls().parse(file_bytes, filename)

        if not doc.raw_text.strip():
            raise ValueError("Document appears to be empty or unreadable after extraction.")

        logger.info(f"Parsed {doc.word_count:,} words, {len(doc.sections)} sections "
                    f"in {doc.parse_time_ms:.0f}ms")
        return doc

    def get_document_stats(self, doc: ParsedDocument) -> Dict:
        return {
            "filename": doc.filename,
            "format": doc.format.upper(),
            "pages": doc.page_count or "N/A",
            "words": f"{doc.word_count:,}",
            "characters": f"{doc.char_count:,}",
            "sections": len(doc.sections),
            "parse_time": f"{doc.parse_time_ms:.0f}ms",
            "metadata": doc.metadata,
        }
